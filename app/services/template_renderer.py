"""模板渲染与 MinIO 存储服务。

使用 python-docx 和 openpyxl 渲染 Word/Excel 模板，
并将产出的文档归档至 MinIO 对象存储。
"""

import hashlib
import io
import logging
import os
from typing import Any

from app.core.config import settings

logger = logging.getLogger(__name__)


class MinioClient:
    """MinIO 对象存储客户端（兼容 S3 协议）。"""

    def __init__(self) -> None:
        self.endpoint = settings.minio_endpoint
        self.access_key = settings.minio_access_key
        self.secret_key = settings.minio_secret_key
        self._client = None

    def _get_client(self):
        """延迟初始化 MinIO 客户端。"""
        if self._client is not None:
            return self._client
        try:
            from minio import Minio
            self._client = Minio(
                endpoint=self.endpoint,
                access_key=self.access_key,
                secret_key=self.secret_key,
                secure=False,
            )
        except ImportError:
            logger.warning("minio 库未安装，使用本地文件系统模拟")
            self._client = None
        return self._client

    def ensure_bucket(self, bucket_name: str) -> None:
        """确保存储桶存在。"""
        client = self._get_client()
        if client is None:
            os.makedirs(f"./local_minio/{bucket_name}", exist_ok=True)
            return
        if not client.bucket_exists(bucket_name):
            client.make_bucket(bucket_name)

    def upload_file(self, bucket: str, object_key: str, data: bytes, content_type: str = "application/octet-stream") -> str:
        """上传文件到 MinIO。

        Returns:
            文件的 MD5 checksum
        """
        checksum = hashlib.md5(data).hexdigest()
        client = self._get_client()
        if client is None:
            local_path = f"./local_minio/{bucket}/{object_key}"
            os.makedirs(os.path.dirname(local_path), exist_ok=True)
            with open(local_path, "wb") as f:
                f.write(data)
            logger.info(f"文件写入本地模拟存储: {local_path}")
            return checksum

        self.ensure_bucket(bucket)
        from minio.commonconfig import Tags
        client.put_object(
            bucket,
            object_key,
            io.BytesIO(data),
            length=len(data),
            content_type=content_type,
        )
        logger.info(f"文件上传 MinIO: {bucket}/{object_key}")
        return checksum

    def download_file(self, bucket: str, object_key: str) -> bytes:
        """从 MinIO 下载文件。"""
        client = self._get_client()
        if client is None:
            local_path = f"./local_minio/{bucket}/{object_key}"
            with open(local_path, "rb") as f:
                return f.read()

        response = client.get_object(bucket, object_key)
        data = response.read()
        response.close()
        response.release_conn()
        return data

    def get_presigned_url(self, bucket: str, object_key: str, expire_seconds: int = 3600) -> str:
        """生成临时下载链接。"""
        client = self._get_client()
        if client is None:
            return f"/local_minio/{bucket}/{object_key}"
        return client.presigned_get_object(bucket, object_key, expires=expire_seconds)


class TemplateRenderer:
    """模板渲染器：将数据填充到 Word/Excel 模板中。"""

    def __init__(self, minio_client: MinioClient | None = None) -> None:
        self.minio = minio_client or MinioClient()

    def _generate_version(self, existing_versions: list[str]) -> str:
        """自动生成版本号（v1, v2, v3...）。"""
        max_ver = 0
        for v in existing_versions:
            v_str = v.lstrip("v")
            if v_str.isdigit():
                max_ver = max(max_ver, int(v_str))
        return f"v{max_ver + 1}"

    def render_word_document(
        self,
        well_data: dict[str, Any],
        measures: list[dict[str, Any]],
        template_path: str | None = None,
    ) -> bytes:
        """渲染 Word 工程设计文档。

        使用 python-docx 将数据填充到 Word 模板中。
        如果模板路径为空，动态生成文档。

        Args:
            well_data: 井基础数据
            measures: 修井措施列表
            template_path: MinIO 模板路径

        Returns:
            渲染后文档的 bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            logger.warning("python-docx 未安装，返回占位数据")
            return b"Engineering design document placeholder"

        if template_path:
            try:
                bucket = settings.minio_bucket_templates
                template_bytes = self.minio.download_file(bucket, template_path)
                doc = Document(io.BytesIO(template_bytes))
            except Exception as exc:
                logger.warning(f"模板加载失败: {exc}，使用空白文档")
                doc = Document()
        else:
            doc = Document()

        # 填充标题
        well_no = well_data.get("well_no", "未知井号")
        title = doc.add_heading(f"井下作业工程设计 — {well_no}", level=1)
        doc.add_paragraph(f"编制时间: {well_data.get('created_at', '')}")

        # 基础信息表
        doc.add_heading("一、基础信息", level=2)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = "Table Grid"
        info_data = [
            ("井号", well_no),
            ("区块", well_data.get("block_name", "")),
            ("层位", well_data.get("layer", "")),
            ("提报单位", well_data.get("report_unit", "")),
            ("产量优先级", str(well_data.get("production_priority", 0))),
            ("上修原因", well_data.get("reason", "")),
        ]
        for i, (key, val) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = str(val)

        # 施工措施
        doc.add_heading("二、修井措施", level=2)
        if measures:
            measures_table = doc.add_table(rows=len(measures) + 1, cols=5)
            measures_table.style = "Table Grid"
            headers = ["措施类型", "工序", "施工参数", "工期(天)", "预估费用"]
            for j, header in enumerate(headers):
                measures_table.rows[0].cells[j].text = header
            for i, m in enumerate(measures):
                measures_table.rows[i + 1].cells[0].text = str(m.get("measure_type", ""))
                measures_table.rows[i + 1].cells[1].text = str(m.get("process", ""))
                measures_table.rows[i + 1].cells[2].text = str(m.get("construction_params", {}))
                measures_table.rows[i + 1].cells[3].text = str(m.get("duration_days", ""))
                measures_table.rows[i + 1].cells[4].text = str(m.get("estimated_cost", ""))

        doc.add_paragraph("")
        doc.add_paragraph("— 文档结束 —")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def render_excel_report(
        self,
        operation_data: list[dict[str, Any]],
        template_path: str | None = None,
    ) -> bytes:
        """渲染 Excel 报表。

        Args:
            operation_data: 作业数据列表
            template_path: MinIO 模板路径

        Returns:
            渲染后文档的 bytes
        """
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, Border, Side
        except ImportError:
            logger.warning("openpyxl 未安装，返回占位数据")
            return b"Excel report placeholder"

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "作业数据报表"

        # 写入标题行
        headers = ["井号", "作业编号", "状态", "进度(%)", "计划开始", "计划结束", "实际开始", "实际结束"]
        for j, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=j, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

        # 写入数据行（列序：井号/作业编号/状态/进度/计划开始/计划结束/实际开始/实际结束）
        for i, record in enumerate(operation_data, 2):
            ws.cell(row=i, column=1, value=record.get("well_no", ""))
            ws.cell(row=i, column=2, value=record.get("operation_no", ""))
            ws.cell(row=i, column=3, value=record.get("status", ""))
            ws.cell(row=i, column=4, value=record.get("progress", 0))
            ws.cell(row=i, column=5, value=str(record.get("planned_start_at", "")))
            ws.cell(row=i, column=6, value=str(record.get("planned_end_at", "")))
            ws.cell(row=i, column=7, value=str(record.get("actual_start_at", "")))
            ws.cell(row=i, column=8, value=str(record.get("actual_end_at", "")))

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def save_to_minio(
        self,
        content: bytes,
        well_no: str,
        version: str,
        bucket: str | None = None,
    ) -> tuple[str, str, str]:
        """将渲染后的文档保存到 MinIO。

        Args:
            content: 文件内容 bytes
            well_no: 井号
            version: 版本号
            bucket: 存储桶名，默认 engineering-designs

        Returns:
            (bucket, object_key, checksum) 元组
        """
        bucket = bucket or settings.minio_bucket_engineering
        object_key = f"engineering/{well_no}/{version}/design.docx"
        self.minio.ensure_bucket(bucket)
        checksum = self.minio.upload_file(bucket, object_key, content)
        return bucket, object_key, checksum
