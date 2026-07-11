from io import BytesIO
from typing import Any

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from sqlalchemy.orm import Session

from app.crud.completion import get_completion_analytics
from app.crud.contractor import get_operation_analytics
from app.crud.material import get_material_analytics
from app.schemas.workover_project_pool import WorkoverAnalyticsQuery
from app.services.statistics_analysis_service import StatisticsAnalysisQuery, build_statistics_analysis
from app.services.workover_analytics_service import build_workover_analytics


def build_delivery_summary(db: Session) -> dict[str, Any]:
    workover = build_workover_analytics(db, WorkoverAnalyticsQuery())
    operations = get_operation_analytics(db)
    materials = get_material_analytics(db)
    completions = get_completion_analytics(db)

    return {
        "projects": {
            "total": workover.kpis.total_projects,
            "pending_approvals": workover.kpis.pending_approvals,
            "approval_rate": workover.kpis.approval_rate,
            "estimated_cost": workover.kpis.estimated_cost,
        },
        "operations": {
            "total_sheets": operations.get("total_sheets", 0),
            "dispatch_rate": operations.get("dispatch_rate", 0),
            "completion_rate": operations.get("completion_rate", 0),
            "team_workload": operations.get("team_workload", []),
        },
        "materials": {
            "total": materials.get("total", 0),
            "delivered": materials.get("delivered", 0),
            "arrived": materials.get("arrived", 0),
            "used": materials.get("used", 0),
            "emergency_count": materials.get("emergency_count", 0),
        },
        "completions": {
            "total": completions.get("total", 0),
            "by_measure_type": completions.get("by_measure_type", []),
        },
    }


def _statistics_payload(db: Session) -> dict[str, Any]:
    try:
        return build_statistics_analysis(db, StatisticsAnalysisQuery())
    except TypeError:
        summary = build_delivery_summary(db)
        return {
            "overview_kpis": {
                "total_projects": summary["projects"]["total"],
                "pending_approvals": summary["projects"]["pending_approvals"],
                "approval_rate": summary["projects"]["approval_rate"],
                "estimated_cost": summary["projects"]["estimated_cost"],
                "operation_sheets": summary["operations"]["total_sheets"],
                "a5_anomalies": 0,
                "material_requirements": summary["materials"]["total"],
                "completion_records": summary["completions"]["total"],
            },
            "operation_efficiency": summary["operations"],
            "a5_statistics": {"anomaly_total": 0, "special_process_total": 0},
            "material_usage": summary["materials"],
            "completion_classification": summary["completions"],
            "trace_sources": ["workover_project_pool", "workover_operation_sheet", "material_requirement", "well_completion_record"],
            "chart_series": {},
        }


def _query_statistics_payload(db: Session, query: StatisticsAnalysisQuery | None = None) -> dict[str, Any]:
    return build_statistics_analysis(db, query or StatisticsAnalysisQuery())


def _append_rows(ws, rows: list[tuple[Any, ...]]) -> None:
    for row in rows:
        ws.append(row)


def _style_header(ws) -> None:
    fill = PatternFill("solid", fgColor="D9EAF7")
    for row in ws.iter_rows(min_row=1, max_row=1):
        for cell in row:
            cell.font = Font(bold=True)
            cell.fill = fill
    for column_cells in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 4, 14), 32)


def export_delivery_summary_excel(db: Session) -> bytes:
    data = _statistics_payload(db)
    overview = data.get("overview_kpis", {})
    operation = data.get("operation_efficiency", {})
    a5 = data.get("a5_statistics", {})
    material = data.get("material_usage", {})
    completion = data.get("completion_classification", {})

    wb = Workbook()
    ws = wb.active
    ws.title = "数据统计分析"
    _append_rows(
        ws,
        [
            ("模块", "指标", "值"),
            ("项目总览", "项目总量", overview.get("total_projects", 0)),
            ("项目总览", "待审批", overview.get("pending_approvals", 0)),
            ("项目总览", "审批通过率", overview.get("approval_rate", 0)),
            ("项目总览", "预计费用(万元)", overview.get("estimated_cost", 0)),
            ("作业运行效率", "运行表工单", operation.get("total_sheets", 0)),
            ("作业运行效率", "派工率", operation.get("dispatch_rate", 0)),
            ("作业运行效率", "完工率", operation.get("completion_rate", 0)),
            ("A5异常与特殊工序", "异常情况", a5.get("anomaly_total", 0)),
            ("A5异常与特殊工序", "特殊工序", a5.get("special_process_total", 0)),
            ("物料使用闭环", "物料需求", material.get("total", 0)),
            ("物料使用闭环", "已到场", material.get("arrived", 0)),
            ("物料使用闭环", "已使用", material.get("used", 0)),
            ("物料使用闭环", "应急需求", material.get("emergency_count", 0)),
            ("完井分类台账", "完井记录", completion.get("total", 0)),
        ],
    )
    _style_header(ws)

    ws_completion = wb.create_sheet("完井分类")
    ws_completion.append(("措施类型", "数量"))
    for item in completion.get("by_measure_type", []):
        ws_completion.append((item.get("measure_type"), item.get("count")))
    _style_header(ws_completion)

    ws_trace = wb.create_sheet("数据追溯")
    ws_trace.append(("追溯来源",))
    for source in data.get("trace_sources", []):
        ws_trace.append((source,))
    _style_header(ws_trace)

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def export_delivery_summary_word(db: Session) -> bytes:
    data = _statistics_payload(db)
    overview = data.get("overview_kpis", {})
    operation = data.get("operation_efficiency", {})
    a5 = data.get("a5_statistics", {})
    material = data.get("material_usage", {})
    completion = data.get("completion_classification", {})

    doc = Document()
    doc.add_heading("采油二厂井下作业管理系统数据统计分析阶段报告", level=1)
    doc.add_paragraph("本报告围绕项目总览、作业运行效率、A5异常与特殊工序、物料使用闭环、完井分类台账形成阶段性统计摘要。")

    doc.add_heading("项目总览", level=2)
    doc.add_paragraph(
        f"项目总量 {overview.get('total_projects', 0)} 项，待审批 {overview.get('pending_approvals', 0)} 项，"
        f"审批通过率 {overview.get('approval_rate', 0)}%，预计费用 {overview.get('estimated_cost', 0)} 万元。"
    )

    doc.add_heading("作业运行效率", level=2)
    doc.add_paragraph(
        f"运行表工单 {operation.get('total_sheets', 0)} 个，派工率 {operation.get('dispatch_rate', 0)}%，"
        f"完工率 {operation.get('completion_rate', 0)}%。"
    )

    doc.add_heading("A5异常与特殊工序", level=2)
    doc.add_paragraph(f"A5异常 {a5.get('anomaly_total', 0)} 条，特殊工序 {a5.get('special_process_total', 0)} 条。")

    doc.add_heading("物料使用闭环", level=2)
    doc.add_paragraph(
        f"物料需求 {material.get('total', 0)} 条，已到场 {material.get('arrived', 0)} 条，"
        f"已使用 {material.get('used', 0)} 条，应急需求 {material.get('emergency_count', 0)} 条。"
    )

    doc.add_heading("完井分类台账", level=2)
    doc.add_paragraph(f"完井记录 {completion.get('total', 0)} 条，按措施类型沉淀分类台账。")

    doc.add_heading("统计结果可追溯", level=2)
    doc.add_paragraph("、".join(data.get("trace_sources", [])) or "暂无追溯来源")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def export_statistics_analysis_excel(db: Session, query: StatisticsAnalysisQuery | None = None) -> bytes:
    data = _query_statistics_payload(db, query)
    overview = data.get("overview_kpis", {})
    operation = data.get("operation_efficiency", {})
    a5 = data.get("a5_statistics", {})
    material = data.get("material_usage", {})
    completion = data.get("completion_classification", {})
    quality = data.get("data_quality_summary", {})

    wb = Workbook()
    ws = wb.active
    ws.title = "数据统计分析"
    _append_rows(
        ws,
        [
            ("模块", "指标", "值"),
            ("项目总览", "项目总量", overview.get("total_projects", 0)),
            ("项目总览", "待审批", overview.get("pending_approvals", 0)),
            ("项目总览", "审批通过率", overview.get("approval_rate", 0)),
            ("项目总览", "预计费用(万元)", overview.get("estimated_cost", 0)),
            ("作业运行效率", "运行表工单", operation.get("total_sheets", 0)),
            ("作业运行效率", "派工率", operation.get("dispatch_rate", 0)),
            ("作业运行效率", "完工率", operation.get("completion_rate", 0)),
            ("A5异常与特殊工序", "异常情况", a5.get("anomaly_total", 0)),
            ("A5异常与特殊工序", "特殊工序", a5.get("special_process_total", 0)),
            ("物料使用闭环", "物料需求", material.get("total", 0)),
            ("物料使用闭环", "已到场", material.get("arrived", 0)),
            ("物料使用闭环", "已使用", material.get("used", 0)),
            ("完井分类台账", "完井记录", completion.get("total", 0)),
            ("数据质量", "问题总数", quality.get("total_issues", 0)),
            ("数据质量", "高风险", quality.get("severity_counts", {}).get("high", 0)),
            ("数据质量", "中风险", quality.get("severity_counts", {}).get("medium", 0)),
            ("数据质量", "低风险", quality.get("severity_counts", {}).get("low", 0)),
        ],
    )
    _style_header(ws)

    ws_quality = wb.create_sheet("数据质量明细")
    ws_quality.append(("规则", "级别", "对象", "井号", "班组", "提示"))
    for issue in quality.get("issues", []):
        ws_quality.append(
            (
                issue.get("title"),
                issue.get("severity"),
                issue.get("entity_type"),
                issue.get("well_no"),
                issue.get("team_name"),
                issue.get("message"),
            )
        )
    _style_header(ws_quality)

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def export_statistics_analysis_word(db: Session, query: StatisticsAnalysisQuery | None = None) -> bytes:
    data = _query_statistics_payload(db, query)
    overview = data.get("overview_kpis", {})
    operation = data.get("operation_efficiency", {})
    a5 = data.get("a5_statistics", {})
    material = data.get("material_usage", {})
    completion = data.get("completion_classification", {})
    quality = data.get("data_quality_summary", {})

    doc = Document()
    doc.add_heading("采油二厂井下作业管理系统数据统计分析报告", level=1)
    doc.add_paragraph(
        f"项目总量 {overview.get('total_projects', 0)}，待审批 {overview.get('pending_approvals', 0)}，"
        f"审批通过率 {overview.get('approval_rate', 0)}%，预计费用 {overview.get('estimated_cost', 0)} 万元。"
    )
    doc.add_heading("作业运行效率", level=2)
    doc.add_paragraph(
        f"运行表工单 {operation.get('total_sheets', 0)}，派工率 {operation.get('dispatch_rate', 0)}%，"
        f"完工率 {operation.get('completion_rate', 0)}%。"
    )
    doc.add_heading("A5异常与特殊工序", level=2)
    doc.add_paragraph(f"A5异常 {a5.get('anomaly_total', 0)} 条，特殊工序 {a5.get('special_process_total', 0)} 条。")
    doc.add_heading("物料使用闭环", level=2)
    doc.add_paragraph(
        f"物料需求 {material.get('total', 0)} 条，已到场 {material.get('arrived', 0)} 条，"
        f"已使用 {material.get('used', 0)} 条。"
    )
    doc.add_heading("完井分类台账", level=2)
    doc.add_paragraph(f"完井记录 {completion.get('total', 0)} 条。")
    doc.add_heading("数据质量", level=2)
    doc.add_paragraph(
        f"问题总数 {quality.get('total_issues', 0)}，"
        f"高风险 {quality.get('severity_counts', {}).get('high', 0)}，"
        f"中风险 {quality.get('severity_counts', {}).get('medium', 0)}，"
        f"低风险 {quality.get('severity_counts', {}).get('low', 0)}。"
    )
    if quality.get("issues"):
        for issue in quality["issues"][:10]:
            doc.add_paragraph(f"{issue.get('title')}: {issue.get('message')}", style="List Bullet")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
